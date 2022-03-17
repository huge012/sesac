from docx import Document

doc = Document() # 논리적 문서 개체 생성 (ms word 실행과 동일)
doc.add_heading("문서 자동화 테스트")   # 본문 제목
doc.add_paragraph("안녕하세요.") # 본문 내용
doc.save("demo1.docx")  # 해당 이름으로 저장
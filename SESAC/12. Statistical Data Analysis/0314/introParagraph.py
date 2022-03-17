from docx import Document
from docx.enum.style import *
from docx.enum.text import *

"""
doc = Document() # 문서 객체 생성
doc.add_paragraph()  # 단락 추가
doc.add_paragraph('두 번째 단락') # 단락 추가

ns = doc.styles["Normal"]
doc.add_paragraph("세번째 단락입니다.")
doc.save("para1.docx")
"""

"""
doc = Document()
for style in doc.styles:
    if style.type == WD_STYLE_TYPE.PARAGRAPH:
        doc.add_paragraph(style.name, style=style)  # 단락의 스타일 전부 보여줌
doc.save('para2.docx')
"""

doc = Document()
para = doc.add_paragraph("New 단락입니다.")  # 단락 객체를 변수로 참조
para.add_run("현재 단락에 run을 추가하였습니다.")
para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
doc.add_paragraph("단락 2")
doc.add_paragraph("단락 3")
p4 = doc.add_paragraph("단락 4")
p4.clear()  # 단락 삭제
para.insert_paragraph_before("단락 앞에 단락을 추가합니다.")
print(para.text)
doc.save("para3.docx")
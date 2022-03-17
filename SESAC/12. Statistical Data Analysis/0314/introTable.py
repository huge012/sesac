from docx import Document
from docx.enum.style import *
from docx.enum.text import *

doc = Document()
""" 스타일 종류 확인
for style in doc.styles:
    if style.type == WD_STYLE_TYPE.TABLE:
        print(style.name)
"""

# Table Grid
from docx.shared import Inches
doc.add_heading("테이블 사용해 봅시다.")
tg_style = doc.styles['Table Grid']
table = doc.add_table(2, 5, tg_style)
doc.save('table1.docx')

for r in range(len(table.rows)):
    row = table.rows[r]
    for c in range(len(row.cells)):
        cell = row.cells[c]
        cell.text = f"{r}, {c}"
doc.save("table2.docx")

cell = table.cell(0, 0) # 테이블의 특정 쉘을 선택
para = cell.add_paragraph() # 쉘에 단락 추가하기
pr = para.add_run() # 단락에 Run 추가하기
pr.add_picture('nasa-Yj1M5riCKk4-unsplash.jpg', width=Inches(1), height=Inches(1))
cell.width = Inches(2.0)
table.rows[0].height = Inches(1.0)
table.rows[1].height = Inches(1.0)
doc.save("table3.docx")

doc = Document()
doc.add_picture('nasa-Yj1M5riCKk4-unsplash.jpg')
doc.save('image.docx')

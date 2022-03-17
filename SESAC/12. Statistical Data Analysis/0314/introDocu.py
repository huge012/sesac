from docx import Document

doc1 = Document()  # 인자 없이 문서 객체 생성
doc1.add_heading("Document1")
doc1.save("doc1.docx")

doc2 = Document("doc1.docx")  # 존재하는 docx 파일 경로를 입력 인자로 전달
doc2.add_paragraph("단락을 추가합니다.")
doc2.save("doc2.docx")  # 존재하는 문서에 새 단락을 추가한 후 다른 이름으로 저장


# 메서드
# add_heading() 헤더 추가
# add_page_break() 페이지 나누기(ctrl+enter)
# add_picture() 사진 추가
# add_table() 테이블 추가
# save() 저장하기
# add_paragraph() 단락추가